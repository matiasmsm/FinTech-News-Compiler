import pinboard
import datetime
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY


pb = pinboard.Pinboard('lsanz_observatorio:CC8E7F9E71CC9B935378')


def escribir_resumen_semanal():
    posts = pb.posts.recent(tag=["pw"])
    print(posts)
    lista_contenido_descripcion = list()
    lista_contenido_url = list()
    with open("Resumenes Semanales/resumen_semanal{}.txt".format(
            datetime.datetime.now()), "w") as resumen_semanal_file:
        for post in posts["posts"]:
            resumen_semanal_file.write(post.description+"\n")
            resumen_semanal_file.write(post.url + "\n \n")
            lista_contenido_descripcion.append(post.description)
            lista_contenido_url.append(post.url)
    escribir_docx(lista_contenido_descripcion, lista_contenido_url)
    escribir_html("Resumen Semanal", lista_contenido_descripcion, lista_contenido_url)


def escribir_pdf(lista_contenido):
    doc = SimpleDocTemplate("Resumenes Semanales/resumen_semanal{}.pdf".format(
            datetime.datetime.now()), pagesize=letter)
    width, height = letter
    Story = []
    logo = "Fotos/logo_bcch.png"
    im = Image(logo, inch, inch)
    Story.append(im)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))
    titulo = "Resumen Semanal Noticias FinTech"
    ptext = '<font size=12>%s</font>'%titulo
    Story.append(Paragraph(ptext, styles["Normal"]))
    Story.append(Spacer(1, 12))
    for linea in lista_contenido:
        ptext = '<font size=12>%s</font>'%linea
        Story.append(Paragraph(ptext, styles["Normal"]))
        if linea[:4] != ["h", "t", "t", "p"]:
            Story.append(Spacer(1, 12))
    doc.build(Story)


def escribir_docx(lista_descripciones, lista_urls):
    document = Document()
    document.add_heading('Resumen Semanal de Noticias FinTech', 0)
    indice_noticia = 0
    for noticia in lista_descripciones:
        document.add_paragraph(noticia)
        document.add_paragraph(lista_urls[indice_noticia])
        indice_noticia += 1
    """document.add_heading('Heading, level 1', level=1)
    document.add_picture('monty-truth.png', width=Inches(1.25))"""
    document.add_page_break()
    document.save("Resumenes Semanales/resumen_semanal{}.docx".format(
            datetime.datetime.now()))


def escribir_html(titulo, lista_descripciones, lista_urls):
    contenido_html="""<html>
<head><title>{}</title></head>
<body>\n""".format(titulo)
    indice_noticia = 0
    for noticia in lista_descripciones:
        contenido_html += """<h2>{}</h2>""".format(noticia)
        contenido_html += """<p>{}</p>""".format(lista_urls[indice_noticia])
        indice_noticia += 1
    contenido_html += """</body>\n</html>"""
    with open("{}.html".format(titulo), "w") as html_file:
        html_file.write(contenido_html)


if __name__ == '__main__':
    escribir_resumen_semanal()
